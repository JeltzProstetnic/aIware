#!/usr/bin/env python3
"""Build two .docx files for Philosophical Psychology submission:
1. Title page (with author info, not sent to reviewers)
2. Anonymized manuscript (sent to reviewers)
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SOURCE = Path("/home/jeltz/aIware/paper/intelligence/paper.md")
OUT_DIR = Path("/home/jeltz/aIware/tmp")

TITLE = "Why Intelligence Models Must Include Motivation: A Recursive Framework"
RUNNING_HEAD = "INTELLIGENCE, MOTIVATION, AND RECURSIVE DEVELOPMENT"

ABSTRACT_TRIMMED = (
    "The field widely acknowledges that motivation influences cognitive development, "
    "yet every major intelligence model formally excludes it. The Cattell-Horn-Carroll "
    "taxonomy contains no motivational component; Cattell\u2019s investment theory treats "
    "motivation as an external condition rather than a constitutive element; Sternberg\u2019s "
    "triarchic theory incorporates practical intelligence but not the drive to acquire it. "
    "This paper argues that this exclusion distorts our understanding of intelligence in "
    "three ways. First, it mischaracterizes intelligence as a static trait rather than a "
    "recursive system in which knowledge, performance, and motivation form a closed "
    "amplification loop. Second, it obscures the role of operational knowledge \u2014 learning "
    "strategies, logical tools, and metacognitive skills \u2014 which functions as the primary "
    "multiplier within this loop. Third, it leaves the field unable to explain why artificial "
    "intelligence systems possessing vast knowledge and computational performance but no "
    "intrinsic motivation fail to exhibit self-directed development. A three-component "
    "recursive model (Knowledge \u00d7 Performance \u00d7 Motivation) is proposed, arguing that "
    "intelligence is best understood not as a capacity but as a learning ability whose "
    "trajectory is determined by recursive dynamics rather than by any single component "
    "measured in isolation."
)

KEYWORDS = "intelligence, motivation, recursive systems, operational knowledge, artificial intelligence, CHC theory, Cattell investment theory"


def set_style(doc):
    """Set default font to Times New Roman 12pt, double spacing."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 2.0
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_running_head(doc, text):
    """Add running head to header."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = text
        p.style.font.name = "Times New Roman"
        p.style.font.size = Pt(12)
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def build_full_with_author():
    """Build complete manuscript WITH author details (non-anonymized)."""
    md_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    set_style(doc)
    add_running_head(doc, RUNNING_HEAD)

    # --- Title page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    run = p.add_run(TITLE)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Matthias Gruber")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Independent Researcher")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("ORCID: 0009-0005-9697-1665")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Correspondence: matthias@matthiasgruber.com")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Word count: approximately 7,800 (body text)")
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Page break after title page
    doc.add_page_break()

    # --- Abstract page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(ABSTRACT_TRIMMED)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Keywords: ")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run = p.add_run(KEYWORDS)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_page_break()

    # --- Body (NOT anonymized) ---
    match = re.search(r'^## 1\. ', md_text, re.MULTILINE)
    if not match:
        raise ValueError("Could not find Section 1")
    body = md_text[match.start():]

    # Do NOT anonymize — keep Gruber citations as-is
    process_body_into_doc(doc, body)

    out = OUT_DIR / "pp-manuscript-with-author.docx"
    doc.save(str(out))
    print(f"Full manuscript (with author): {out}")


def parse_body(md_text):
    """Parse the markdown body into structured elements, skipping front matter."""
    # Remove front matter (title, author, target journal, abstract, keywords)
    # Find where Section 1 starts
    match = re.search(r'^## 1\. ', md_text, re.MULTILINE)
    if not match:
        raise ValueError("Could not find Section 1")
    body = md_text[match.start():]

    # Also grab references section
    return body


def anonymize_text(text):
    """Replace self-identifying references."""
    # Gruber (2015) -> [Author] (2015)
    text = re.sub(r'Gruber\s*\(2015\)', '[Author] (2015)', text)
    # Gruber (2026) -> [Author] (2026)
    text = re.sub(r'Gruber\s*\(2026\)', '[Author] (2026)', text)
    # Gruber, submitted -> [Author], submitted
    text = re.sub(r'Gruber,\s*submitted', '[Author], submitted', text)
    # "first proposed in Gruber (2015, in German)"
    text = re.sub(r'Gruber\s*\(2015,\s*in German\)', '[Author] (2015, in German)', text)
    # In references section
    text = re.sub(r'^Gruber, M\. \(2015\)\. \*Die Emergenz des Bewusstseins\*\. Self-published\. ISBN 9781326652074\.',
                  '[Author]. (2015). [Title withheld for blind review]. Self-published.',
                  text, flags=re.MULTILINE)
    text = re.sub(r'^Gruber, M\. \(2026\)\. The four-model theory.*$',
                  '[Author]. (2026). [Title withheld for blind review]. Zenodo preprint.',
                  text, flags=re.MULTILINE)
    return text


def add_formatted_paragraph(doc, text, is_heading=False, heading_level=1, is_italic=False):
    """Add a paragraph with inline formatting (bold, italic)."""
    if is_heading:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        if heading_level == 1:
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif heading_level == 2:
            run.bold = True
            run.italic = True
        elif heading_level == 3:
            run.italic = True
        return p

    p = doc.add_paragraph()

    # Parse inline markdown formatting
    # Handle *text* (italic) and **text** (bold)
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    return p


def process_body_into_doc(doc, body):
    """Process markdown body text into a docx document."""
    lines = body.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        if line.strip() == '---':
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith('## '):
            heading_text = re.sub(r'^##\s+', '', line)
            add_formatted_paragraph(doc, heading_text, is_heading=True, heading_level=1)
            i += 1
            continue

        if line.startswith('### '):
            heading_text = re.sub(r'^###\s+', '', line)
            add_formatted_paragraph(doc, heading_text, is_heading=True, heading_level=2)
            i += 1
            continue

        if re.match(r'^\d+\.\s+', line):
            text = line
            while i + 1 < len(lines) and lines[i + 1].strip() and not re.match(r'^\d+\.\s+', lines[i + 1]) and not lines[i + 1].startswith('#') and not lines[i + 1].startswith('-') and lines[i + 1].strip() != '---':
                i += 1
                text += ' ' + lines[i].strip()
            add_formatted_paragraph(doc, text)
            i += 1
            continue

        if line.startswith('- ') or line.startswith('   - '):
            text = line.lstrip('- ').strip()
            while i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].startswith('-') and not lines[i + 1].startswith('#') and not re.match(r'^\d+\.\s+', lines[i + 1]) and lines[i + 1].strip() != '---':
                i += 1
                text += ' ' + lines[i].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    run = p.add_run(part)
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            i += 1
            continue

        text = line
        while i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].startswith('#') and not lines[i + 1].startswith('-') and not re.match(r'^\d+\.\s+', lines[i + 1]) and lines[i + 1].strip() != '---':
            i += 1
            text += ' ' + lines[i].strip()

        add_formatted_paragraph(doc, text)
        i += 1


def anonymize_body_for_blind(body):
    """Remove Acknowledgments section from body for blind review, keep Disclosure Statement."""
    # Replace "## Acknowledgments" with "## Disclosure Statement" and strip
    # identifying content, keeping only the AI disclosure
    body = re.sub(
        r'^## Acknowledgments\n\n.*?(?=\n---|\n## )',
        '## Disclosure Statement\n\n'
        'The author used Claude (Anthropic, Claude Opus 4) for editorial assistance '
        'and manuscript formatting. All theoretical content, arguments, and conclusions '
        'are solely the author\u2019s own.\n\n',
        body, flags=re.MULTILINE | re.DOTALL
    )
    return body


def build_manuscript_anonymous():
    """Build anonymized manuscript (no author info, anonymized self-citations)."""
    md_text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    set_style(doc)
    add_running_head(doc, RUNNING_HEAD)

    # Title (no author)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    run = p.add_run(TITLE)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    # Abstract
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("Abstract")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run(ABSTRACT_TRIMMED)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Keywords: ")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run = p.add_run(KEYWORDS)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

    doc.add_page_break()

    match = re.search(r'^## 1\. ', md_text, re.MULTILINE)
    if not match:
        raise ValueError("Could not find Section 1")
    body = md_text[match.start():]
    body = anonymize_text(body)
    body = anonymize_body_for_blind(body)

    process_body_into_doc(doc, body)

    out = OUT_DIR / "pp-manuscript-anonymous.docx"
    doc.save(str(out))
    print(f"Anonymous manuscript saved: {out}")

    # Word count
    ref_match = re.search(r'^## References', body, re.MULTILINE)
    body_only = body[:ref_match.start()] if ref_match else body
    body_only = re.sub(r'[#*\-]', '', body_only)
    print(f"Body word count (approx): {len(body_only.split())}")


if __name__ == "__main__":
    build_full_with_author()
    build_manuscript_anonymous()
    print("Done!")
